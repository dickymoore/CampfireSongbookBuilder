from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import logging
from app.fetch_data import get_lyrics_from_genius
from app.cache import cache_data, load_cache

# Configure logging
logger = logging.getLogger(__name__)

def remove_contributors_and_embeds(lyrics):
    """Remove the Contributors, Embed sections, and unwanted advertisements from the lyrics."""
    lyrics = re.sub(r'^.*Contributors', '', lyrics, flags=re.DOTALL)
    lyrics = re.sub(r'Embed\s*$', '', lyrics, flags=re.MULTILINE)
    return lyrics

def remove_unwanted_phrases(lyrics):
    """Remove unwanted phrases and advertisements from the lyrics without removing the entire line."""
    patterns_to_remove = [
        r'You might also like.*?',
        r'See .*? LiveGet tickets as low as \$\d+',
        r'.*? Lyrics'
    ]
    
    # Remove each pattern, but only the matched part, not the entire line
    for pattern in patterns_to_remove:
        lyrics = re.sub(pattern, '', lyrics, flags=re.MULTILINE)
    
    return lyrics

def clean_lyrics(lyrics):
    """Clean the lyrics by removing contributors, embeds, and unwanted phrases."""
    lyrics = remove_contributors_and_embeds(lyrics)
    lyrics = remove_unwanted_phrases(lyrics)
    return lyrics

def set_document_margins(document, margin_in_inches):
    """Set the margins of the document."""
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(margin_in_inches)
        section.bottom_margin = Inches(margin_in_inches)
        section.left_margin = Inches(margin_in_inches)
        section.right_margin = Inches(margin_in_inches)

def set_paragraph_font(paragraph, font_size):
    """Set the font size of a paragraph."""
    for run in paragraph.runs:
        run.font.size = Pt(font_size)

def create_two_column_section(document):
    """Create a new section with two columns."""
    section = document.add_section()
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

def add_header_footer(document):
    """Add a header and footer with page numbers to the document."""
    # Add header
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = "Campfire Songs"
    paragraph.style.font.size = Pt(14)

    # Add footer with page numbers
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Page "
    paragraph.style.font.size = Pt(12)
    
    # Add the page number field to the footer
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def sort_songs(song_list):
    """Sort songs case-insensitively and ignoring special characters."""
    return sorted(song_list, key=lambda x: re.sub(r'[^a-zA-Z0-9]', '', x['Title']).lower())

def get_song_lyrics_info(song_list, genius_client):
    """Get song titles and the character length of the lyrics for those songs."""
    lyrics_cache = load_cache('data/cache/lyrics_cache.json')
    song_info = []

    sorted_songs = sort_songs(song_list)

    for song in sorted_songs:
        artist = song['Artist']
        title = song['Title']
        cache_key = f"{artist} - {title}"
        
        if cache_key in lyrics_cache and bool(lyrics_cache[cache_key]) and lyrics_cache[cache_key] != "Lyrics not found.":
            lyrics = lyrics_cache[cache_key]
            logger.debug("Lyrics loaded from cache.")
        else:
            lyrics = get_lyrics_from_genius(title, artist, genius_client)
            cleaned_lyrics = clean_lyrics(lyrics)
            num_characters = len(cleaned_lyrics)
            
            if bool(lyrics) and lyrics != "Lyrics not found." and num_characters <= 5000:
                lyrics_cache[cache_key] = lyrics
                cache_data('data/cache/lyrics_cache.json', lyrics_cache)  # Update the cache file immediately
                logger.debug("Lyrics fetched and cached.")
            else:
                lyrics = "Lyrics not found."
                logger.debug("Lyrics not found or too long.")

        cleaned_lyrics = clean_lyrics(lyrics)
        num_characters = len(cleaned_lyrics)
        song_info.append((title, num_characters))

    return song_info

def generate_documents(song_list, genius_client, lyrics_output, chords_output):
    logger.info("Loading cache data...")
    lyrics_cache = load_cache('data/cache/lyrics_cache.json')
    chords_cache = load_cache('data/cache/chords_cache.json')

    lyrics_document = Document()
    chords_document = Document()

    # Set document margins to 0.5 inches
    set_document_margins(lyrics_document, 0.5)
    set_document_margins(chords_document, 0.5)

    # Create two-column section
    create_two_column_section(lyrics_document)

    # Add header and footer with page numbers
    add_header_footer(lyrics_document)

    sorted_songs = sort_songs(song_list)

    for song in sorted_songs:
        artist = song['Artist']
        title = song['Title']
        logger.debug(f"Processing {title} by {artist}...")

        cache_key = f"{artist} - {title}"
        if cache_key in lyrics_cache and bool(lyrics_cache[cache_key]):
            lyrics = lyrics_cache[cache_key]
            logger.debug("Lyrics loaded from cache.")
        else:
            lyrics = get_lyrics_from_genius(title, artist, genius_client)
            cleaned_lyrics = clean_lyrics(lyrics)
            num_characters = len(cleaned_lyrics)
            
            if lyrics and lyrics != "Lyrics not found." and num_characters <= 5000:
                lyrics_cache[cache_key] = lyrics
                cache_data('data/cache/lyrics_cache.json', lyrics_cache)  # Update the cache file immediately
                logger.debug("Lyrics fetched and cached.")
            else:
                lyrics = "Lyrics not found."
                logger.debug("Lyrics not found or too long.")

        cleaned_lyrics = clean_lyrics(lyrics)
        num_characters = len(cleaned_lyrics)
        
        if num_characters <= 5000:
            # Add heading
            heading = lyrics_document.add_heading(f"{title} by {artist}", level=1)
            set_paragraph_font(heading, 14)  # Set heading font size

            # Add lyrics
            paragraph = lyrics_document.add_paragraph(cleaned_lyrics)
            set_paragraph_font(paragraph, 12)  # Set lyrics font size
        else:
            logger.debug(f"Lyrics for {title} are too long and have been excluded.")

    lyrics_document.save(lyrics_output)
    logger.info(f"Lyrics document saved as {lyrics_output}.")

    # Save chords document if needed
    # chords_document.save(chords_output)
    # logger.info(f"Chords document saved as {chords_output}.")
