from docx import Document
import logging
from app.fetch_data import get_lyrics_from_genius, get_chords_from_chordie
from app.cache import cache_data, load_cache
from app.text_cleaning import clean_lyrics
from app.document_formatting import set_document_margins, set_paragraph_font, create_two_column_section, add_header_footer, sort_songs

# Configure logging
logger = logging.getLogger(__name__)

def get_song_lyrics_info(song_list, genius_client):
    """Get song titles and the character length of the lyrics for those songs."""
    lyrics_cache = load_cache('data/cache/lyrics_cache.json')
    song_info = []

    sorted_songs = sort_songs(song_list)

    for song in sorted_songs:
        artist = song['Artist']
        title = song['Title']
        cache_key = f"{artist} - {title}"
        
        if cache_key in lyrics_cache and bool(lyrics_cache[cache_key]) and lyrics_cache[cache_key] != "Lyrics not found.":
            lyrics = lyrics_cache[cache_key]
            logger.debug("Lyrics loaded from cache.")
        else:
            lyrics = get_lyrics_from_genius(title, artist, genius_client)
            cleaned_lyrics = clean_lyrics(lyrics)
            num_characters = len(cleaned_lyrics)
            
            if bool(lyrics) and lyrics != "Lyrics not found." and num_characters <= 5000:
                lyrics_cache[cache_key] = lyrics
                cache_data('data/cache/lyrics_cache.json', lyrics_cache)  # Update the cache file immediately
                logger.debug("Lyrics fetched and cached.")
            else:
                lyrics = "Lyrics not found."
                logger.debug("Lyrics not found or too long.")

        cleaned_lyrics = clean_lyrics(lyrics)
        num_characters = len(cleaned_lyrics)
        song_info.append((title, num_characters))

    return song_info

def generate_documents(song_list, genius_client, lyrics_output, chords_output):
    logger.info("Loading cache data...")
    lyrics_cache = load_cache('data/cache/lyrics_cache.json')
    chords_cache = load_cache('data/cache/chords_cache.json')

    lyrics_document = Document()
    chords_document = Document()

    # Set document margins to 0.5 inches
    set_document_margins(lyrics_document, 0.5)
    set_document_margins(chords_document, 0.5)

    # Create two-column section
    create_two_column_section(lyrics_document)

    # Add header and footer with page numbers
    add_header_footer(lyrics_document)
    add_header_footer(chords_document)

    sorted_songs = sort_songs(song_list)

    for song in sorted_songs:
        artist = song['Artist']
        title = song['Title']
        logger.debug(f"Processing {title} by {artist}...")

        cache_key = f"{artist} - {title}"
        if cache_key in lyrics_cache and bool(lyrics_cache[cache_key]):
            lyrics = lyrics_cache[cache_key]
            logger.debug("Lyrics loaded from cache.")
        else:
            logger.debug(f"Fetching lyrics for {title} by {artist} from Genius API...")
            lyrics = get_lyrics_from_genius(title, artist, genius_client)
            cleaned_lyrics = clean_lyrics(lyrics)
            num_characters = len(cleaned_lyrics)
            
            if lyrics and lyrics != "Lyrics not found." and num_characters <= 5000:
                lyrics_cache[cache_key] = lyrics
                cache_data('data/cache/lyrics_cache.json', lyrics_cache)
                logger.debug("Lyrics fetched and cached.")
            else:
                lyrics = "Lyrics not found."
                logger.debug("Lyrics not found or too long.")

        cleaned_lyrics = clean_lyrics(lyrics)
        num_characters = len(cleaned_lyrics)
        
        if num_characters <= 5000:
            # Add heading
            heading = lyrics_document.add_heading(f"{title} by {artist}", level=1)
            set_paragraph_font(heading, 14)

            # Add lyrics
            paragraph = lyrics_document.add_paragraph(cleaned_lyrics)
            set_paragraph_font(paragraph, 12)
        else:
            logger.debug(f"Lyrics for {title} are too long and have been excluded.")

        # Fetch and add chords
        if cache_key in chords_cache and bool(chords_cache[cache_key]):
            chords = chords_cache[cache_key]
            logger.debug("Chords loaded from cache.")
        else:
            logger.debug(f"Fetching chords for {title} by {artist} from Chordie...")
            chords = get_chords_from_chordie(title, artist)
            if chords and chords != "Chords not found.":
                chords_cache[cache_key] = chords
                cache_data('data/cache/chords_cache.json', chords_cache)
                logger.debug("Chords fetched and cached.")
            else:
                chords = "Chords not found."
                logger.debug("Chords not found.")

        if chords != "Chords not found.":
            # Add heading
            heading = chords_document.add_heading(f"{title} by {artist}", level=1)
            set_paragraph_font(heading, 14)

            # Add chords
            paragraph = chords_document.add_paragraph(chords)
            set_paragraph_font(paragraph, 12)
        else:
            logger.debug(f"Chords for {title} are not available.")

    lyrics_document.save(lyrics_output)
    logger.info(f"Lyrics document saved as {lyrics_output}.")

    chords_document.save(chords_output)
    logger.info(f"Chords document saved as {chords_output}.")
