import requests
from bs4 import BeautifulSoup

def get_chords_from_ultimate_guitar(song_title, artist_name):
    search_url = "https://www.ultimate-guitar.com/search.php?search_type=title&value= "
    response = requests.get(search_url)
    soup = BeautifulSoup(response.text, 'html.parser')
    result = soup.find("", class_="js-tp_link")
    if result:
        chords_url = result['href']
        chords_response = requests.get(chords_url)
        chords_soup = BeautifulSoup(chords_response.text, 'html.parser')
        chords_div = chords_soup.find("pre", class_="js-tab-content")
        chords = chords_div.get_text() if chords_div else "Chords not found."
        return chords
    return "Chords not found."

def get_lyrics_from_genius(song_title, artist_name, api_key):
    base_url = "https://api.genius.com"
    headers = {'Authorization': f'Bearer '}
    search_url = f""/search""
    params = {'q': f"" ""}
    response = requests.get(search_url, headers=headers, params=params)
    json_response = response.json()
    remote_song_info = None
    for hit in json_response['response']['hits']:
        if artist_name.lower() in hit['result']['primary_artist']['name'].lower():
            remote_song_info = hit
            break
    if remote_song_info:
        song_path = remote_song_info['result']['path']
        song_url = f""""
        song_page = requests.get(song_url)
        song_page_soup = BeautifulSoup(song_page.text, 'html.parser')
        lyrics_div = song_page_soup.find("div", class_="lyrics")
        lyrics = lyrics_div.get_text() if lyrics_div else "Lyrics not found."
        return lyrics
    return "Lyrics not found."
"@

 = @"
from docx import Document
import time
from .fetch_data import get_lyrics_from_genius, get_chords_from_ultimate_guitar
from .cache import cache_data, load_cache

def generate_documents(song_list, api_key, lyrics_output, chords_output):
    lyrics_cache = load_cache('data/lyrics_cache.json')
    chords_cache = load_cache('data/chords_cache.json')

    lyrics_document = Document()
    chords_document = Document()

    for song in sorted(song_list, key=lambda x: x['title']):
        artist = song['artist']
        title = song['title']

        cache_key = f"" - ""
        if cache_key in lyrics_cache:
            lyrics = lyrics_cache[cache_key]
        else:
            lyrics = get_lyrics_from_genius(title, artist, api_key)
            lyrics_cache[cache_key] = lyrics

        if cache_key in chords_cache:
            chords = chords_cache[cache_key]
        else:
            chords = get_chords_from_ultimate_guitar(title, artist)
            chords_cache[cache_key] = chords

        lyrics_document.add_heading(f"" by "", level=1)
        lyrics_document.add_paragraph(lyrics)
        lyrics_document.add_page_break()

        chords_document.add_heading(f"" by "", level=1)
        chords_document.add_paragraph(chords)
        chords_document.add_page_break()

        time.sleep(5)

    lyrics_document.save(lyrics_output)
    chords_document.save(chords_output)

    cache_data('data/lyrics_cache.json', lyrics_cache)
    cache_data('data/chords_cache.json', chords_cache)
