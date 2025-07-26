"""Functions for generating songbook documents from cached data.

The ``create_document_from_cache`` function accepts a list of songs and
pre‑loaded cache dictionaries for lyrics and chords.  It produces Word
documents containing the selected content.  Songs are sorted alphabetically
before being written.
"""

from docx import Document
import logging
from .document_formatting import (
    set_document_margins,
    set_paragraph_font,
    create_two_column_section,
    add_header_footer,
    sort_songs,
)
from .text_cleaning import clean_lyrics, clean_chords

logger = logging.getLogger(__name__)

def create_document_from_cache(
    song_list,
    lyrics_cache: dict,
    chords_cache: dict,
    lyrics_output: str | None = None,
    chords_output: str | None = None,
) -> None:
    """Generate Word documents for lyrics and/or chords from cache.

    Args:
        song_list: List of song dictionaries containing at least 'Artist' and 'Title'.
        lyrics_cache: Dictionary mapping ``"Artist - Title"`` to raw lyrics.
        chords_cache: Dictionary mapping ``"Artist - Title"`` to raw chords.
        lyrics_output: Path to save the lyrics document.  If ``None`` the
            lyrics document is not generated.
        chords_output: Path to save the chords document.  If ``None`` the
            chords document is not generated.

    The function skips entries where the lyrics are missing or exceed 5,000
    characters to avoid very long documents.  Chords are always included
    provided they are present in the cache.
    """
    logger.debug("Running create_document_from_cache function")

    lyrics_document = None
    chords_document = None
    if lyrics_output:
        logger.debug("Initializing lyrics document")
        lyrics_document = Document()
        set_document_margins(lyrics_document, 0.5)
        create_two_column_section(lyrics_document)
        add_header_footer(lyrics_document)
    if chords_output:
        logger.debug("Initializing chords document")
        chords_document = Document()
        set_document_margins(chords_document, 0.5)
        create_two_column_section(chords_document)
        add_header_footer(chords_document)

    sorted_songs = sort_songs(song_list)
    for song in sorted_songs:
        artist = song['Artist']
        title = song['Title']
        cache_key = f"{artist} - {title}"
        # Lyrics section
        if lyrics_document and cache_key in lyrics_cache and bool(lyrics_cache[cache_key]):
            lyrics_raw = lyrics_cache[cache_key]
            lyrics = clean_lyrics(lyrics_raw)
            num_characters = len(lyrics)
            logger.debug("Adding lyrics for %s by %s", title, artist)
            if num_characters <= 5000:
                heading = lyrics_document.add_heading(f"{title} by {artist}", level=1)
                set_paragraph_font(heading, 14)
                paragraph = lyrics_document.add_paragraph()
                lines = lyrics.split('\n')
                for i, line in enumerate(lines):
                    if i > 0:
                        paragraph.add_run().add_break()
                    paragraph.add_run(line)
                set_paragraph_font(paragraph, 12)
            else:
                logger.debug("Lyrics for %s are too long and have been excluded.", title)
        # Chords section
        if chords_document and cache_key in chords_cache and bool(chords_cache[cache_key]):
            chords_raw = chords_cache[cache_key]
            chords = clean_chords(chords_raw)
            logger.debug("Adding chords for %s by %s", title, artist)
            heading = chords_document.add_heading(f"{title} by {artist}", level=1)
            set_paragraph_font(heading, 14)
            paragraph = chords_document.add_paragraph()
            lines = chords.split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    paragraph.add_run().add_break()
                paragraph.add_run(line)
            set_paragraph_font(paragraph, 12)
    # Save documents
    if lyrics_document:
        lyrics_document.save(lyrics_output)
        logger.info("Lyrics document saved as %s.", lyrics_output)
    if chords_document:
        chords_document.save(chords_output)
        logger.info("Chords document saved as %s.", chords_output)