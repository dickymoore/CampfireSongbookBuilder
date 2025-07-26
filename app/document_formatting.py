"""Document formatting helpers.

This module provides utility functions for formatting Word documents using
``python-docx``.  It sets margins, creates multi‑column layouts and adds
headers and footers with page numbers.
"""

# Import optional dependencies from python-docx.  These imports are deferred
# because ``python-docx`` may not be installed in all environments (for
# example, during continuous integration or when only sorting functions are
# needed).  Functions that rely on these imports should check for ``None``
# and raise a descriptive error if necessary.
try:
    from docx.shared import Pt, Inches  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
except ImportError:  # pragma: no cover
    Pt = None  # type: ignore
    Inches = None  # type: ignore
    qn = None  # type: ignore
    OxmlElement = None  # type: ignore
import re

def set_document_margins(document, margin_in_inches: float) -> None:
    """Set uniform margins for all sections in the document.

    Args:
        document: A ``docx.Document`` instance.
        margin_in_inches: The margin size in inches for all sides.
    """
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(margin_in_inches)
        section.bottom_margin = Inches(margin_in_inches)
        section.left_margin = Inches(margin_in_inches)
        section.right_margin = Inches(margin_in_inches)

def set_paragraph_font(paragraph, font_size: int) -> None:
    """Set the font size for all runs in a paragraph.

    Args:
        paragraph: A ``docx.Paragraph`` instance.
        font_size: Font size in points.
    """
    for run in paragraph.runs:
        run.font.size = Pt(font_size)

def create_two_column_section(document) -> None:
    """Add a new section with two columns to the document."""
    section = document.add_section()
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

def add_header_footer(document) -> None:
    """Add a header and footer with page numbers to the first section."""
    # Header
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = "Campfire Songs"
    paragraph.style.font.size = Pt(14)

    # Footer with page numbers
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Page "
    paragraph.style.font.size = Pt(12)
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def sort_songs(song_list):
    """Sort a list of song dictionaries case‑insensitively and ignore special characters.

    Args:
        song_list: A list of dictionaries with at least the key ``'Title'``.

    Returns:
        The sorted list.
    """
    return sorted(song_list, key=lambda x: re.sub(r'[^a-zA-Z0-9]', '', x['Title']).lower())